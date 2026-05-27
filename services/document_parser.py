"""文档解析 — .txt 纯文本 / .docx Word表格（优先提取表格）"""

import os
from config import MAX_CONTENT_LENGTH


def parse_document(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    content = _parse_txt(filepath) if ext == '.txt' else (
        _parse_docx(filepath) if ext == '.docx' else ''
    )
    if not content:
        raise ValueError(f"不支持的文件格式: {ext}")
    return content[:MAX_CONTENT_LENGTH]  # 截断保护


def _parse_txt(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read(MAX_CONTENT_LENGTH)


def _parse_docx(filepath):
    """优先提取表格（分镜脚本常以Word表格存在），无表格则读段落"""
    from docx import Document
    doc = Document(filepath)

    tables_text = []
    for table in doc.tables:
        for row in table.rows:
            tables_text.append(' | '.join(c.text.strip() for c in row.cells))
    if tables_text:
        return '\n'.join(tables_text)

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return '\n'.join(paragraphs[:500])
